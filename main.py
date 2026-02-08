# main.py — AutoRTC Studio (SaaS UI + FULL data-driven report + AI narrative + Home images)
import os
import io
import json
import datetime
from typing import Dict, Any, List, Optional, Tuple

import pandas as pd
import streamlit as st

from utils import read_input_streamlit, preprocess
import descriptive, predictive

from dotenv import load_dotenv
from langchain_core.prompts import PromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_groq import ChatGroq

from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.styles import getSampleStyleSheet

from docx import Document
from docx.shared import Pt

import folium
from streamlit_folium import st_folium


# ---------------- App config ----------------
st.set_page_config(page_title="AutoRTC Studio", page_icon="💥🚗", layout="wide")
load_dotenv()

if "groq_api_key" not in st.session_state:
    st.session_state.groq_api_key = os.getenv("GROQ_API_KEY", "")
if "page" not in st.session_state:
    st.session_state.page = "Home"
if "latest_report" not in st.session_state:
    st.session_state.latest_report = None


# ---------------- SaaS UI ----------------
st.markdown(r"""
<style>
@import url('https://fonts.googleapis.com/css2?family=Plus+Jakarta+Sans:wght@600;700;800&display=swap');
html, body, [class*="css"] { font-family: 'Plus Jakarta Sans', sans-serif; }

:root{
  --bg:#f4f7fa;
  --ink:#0f172a;
  --muted:#64748b;
  --border: rgba(15,23,42,0.08);

  --blue:#00A1DE;
  --yellow:#FAD201;
  --green:#20603D;
  --dark:#0f172a;

  --pink:#E91E63;   /* PDF */
  --indigo:#3F51B5; /* Word */
  --slate:#334155;  /* JSON */
}

.stApp{ background: var(--bg); }
.block-container{ max-width: 1400px; padding-top: 1rem; padding-bottom: 2rem; }

/* Big text */
h1{ font-size: 44px !important; font-weight: 900 !important; letter-spacing: -1.2px; color: var(--ink); }
h2{ font-size: 28px !important; font-weight: 900 !important; color: var(--ink); }
h3{ font-size: 20px !important; font-weight: 900 !important; color: var(--ink); }
p, li, span, div{ font-size: 16px; }

/* Shared cards */
.hero, .section{
  background:#fff;
  border-radius: 24px;
  border: 1px solid var(--border);
  box-shadow: 0 10px 30px rgba(0,0,0,0.03);
  padding: 22px;
  margin-bottom: 18px;
}

/* Sidebar */
section[data-testid="stSidebar"]{ background:#fff !important; border-right: 1px solid rgba(15,23,42,0.06); }
div[role="radiogroup"] label p{ font-size: 18px !important; font-weight: 800 !important; color: var(--ink) !important; }

/* Quick launch cards */
.launch{
  display:flex; flex-direction:column; height:100%;
  border-radius: 22px;
  padding: 20px;
  transition: all .25s ease;
  border: 1px solid rgba(15,23,42,0.06);
}
.launch:hover{
  transform: translateY(-8px);
  box-shadow: 0 20px 40px rgba(0,0,0,0.10);
}
.launch-title{ font-size: 18px; font-weight: 900; margin-bottom: 6px; }
.launch-desc{ font-size: 14px; font-weight: 700; opacity: .85; line-height: 1.4; margin-bottom: 14px; flex-grow: 1; }

/* Card colors */
.card-blue{   background:#eef7ff; color:#0061d5; border-bottom: 4px solid var(--blue); }
.card-yellow{ background:#fffdf2; color:#7a5c00; border-bottom: 4px solid var(--yellow); }
.card-green{  background:#f2faf5; color:#155724; border-bottom: 4px solid var(--green); }
.card-dark{   background:#f8fafc; color:#0f172a; border-bottom: 4px solid var(--dark); }

/* Button base */
div.stButton > button, div.stDownloadButton > button{
  border-radius: 12px !important;
  padding: 12px 18px !important;
  font-weight: 900 !important;
  border: none !important;
  width: 100% !important;
  transition: all .18s ease !important;
}
div.stButton > button:hover, div.stDownloadButton > button:hover{
  transform: translateY(-2px) !important;
  box-shadow: 0 10px 18px rgba(0,0,0,0.08) !important;
}

/* Buttons match cards (wrap the widget in a div) */
.btn-blue   div.stButton > button{ background: var(--blue) !important; color:#fff !important; }
.btn-yellow div.stButton > button{ background: var(--yellow) !important; color:#3a2b00 !important; }
.btn-green  div.stButton > button{ background: var(--green) !important; color:#fff !important; }
.btn-dark   div.stButton > button{ background: var(--dark) !important; color:#fff !important; }

/* Export download button colors */
.json-btn  div.stDownloadButton > button{ background: var(--slate) !important; color:#fff !important; }
.pdf-btn   div.stDownloadButton > button{ background: var(--pink) !important; color:#fff !important; }
.word-btn  div.stDownloadButton > button{ background: var(--indigo) !important; color:#fff !important; }

/* AI panels */
.ai-agent{
  background: rgba(0,161,222,0.06) !important;
  border: 1.5px solid rgba(0,161,222,0.20);
  border-radius: 20px;
  padding: 18px;
  height: 100%;
}
.ai-chat{
  background: rgba(100,116,139,0.08) !important;
  border: 1.5px solid rgba(100,116,139,0.20);
  border-radius: 20px;
  padding: 18px;
  height: 100%;
}

/* Image cards */
.imgcard{
  background: #fff;
  border-radius: 22px;
  border: 1px solid rgba(15,23,42,0.08);
  box-shadow: 0 10px 26px rgba(0,0,0,0.04);
  overflow: hidden;
}
.imgcap{ padding: 12px 14px; font-weight: 900; color: var(--ink); }
.imgsub{ padding: 0 14px 12px 14px; color: var(--muted); font-weight: 700; font-size: 14px; }

/* Uploader */
div[data-testid="stFileUploaderDropzone"]{
  border-radius: 16px !important;
  border: 2px dashed var(--blue) !important;
  background: rgba(255,255,255,0.75) !important;
}
</style>
""", unsafe_allow_html=True)




# ---------------- Helpers ----------------
def safe_int(x, default=0):
    try:
        return int(float(x))
    except Exception:
        return default

def _lower_cols(df: pd.DataFrame) -> Dict[str, str]:
    return {str(c).lower(): c for c in df.columns}

def pick_numeric_sum(df: pd.DataFrame, candidates: List[str]) -> Optional[int]:
    low = _lower_cols(df)
    for cand in candidates:
        if cand in low:
            col = low[cand]
            if pd.api.types.is_numeric_dtype(df[col]):
                return safe_int(df[col].fillna(0).sum(), None)
    return None

def filter_df(df: pd.DataFrame, filters: dict) -> pd.DataFrame:
    out = df.copy()
    if "location_province" in out.columns and filters.get("location_province"):
        out = out[out["location_province"].isin(filters["location_province"])]
    if "location_name" in out.columns and filters.get("location_name"):
        out = out[out["location_name"].isin(filters["location_name"])]
    return out

def detect_time_col(df: pd.DataFrame) -> Optional[str]:
    candidates = ["date", "datetime", "timestamp", "crash_date", "accident_date", "time", "crash_time", "event_date"]
    low = _lower_cols(df)
    for cand in candidates:
        for k in list(low.keys()):
            if cand == k or cand in k:
                return low[k]
    return None

def detect_lat_lon(df: pd.DataFrame) -> Tuple[Optional[str], Optional[str]]:
    lat = lon = None
    for c in df.columns:
        cl = str(c).lower()
        if lat is None and cl in ["lat", "latitude", "gps_lat", "location_lat", "y"]:
            lat = c
        if lon is None and cl in ["lon", "lng", "longitude", "gps_lon", "location_lon", "x"]:
            lon = c
    return lat, lon

def _safe_value_counts(df: pd.DataFrame, col: str, n: int = 10) -> Dict[str, int]:
    if col not in df.columns:
        return {}
    s = df[col].dropna()
    if s.empty:
        return {}
    return s.astype(str).value_counts().head(n).to_dict()

def _missingness(df: pd.DataFrame, top_n: int = 12) -> List[Dict[str, Any]]:
    miss = (df.isna().sum() / max(len(df), 1) * 100).sort_values(ascending=False)
    out = []
    for col, pct in miss.head(top_n).items():
        out.append({"column": str(col), "missing_pct": round(float(pct), 2), "missing_count": int(df[col].isna().sum())})
    return out

def _time_insights(df: pd.DataFrame, time_col: str) -> Dict[str, Any]:
    s = pd.to_datetime(df[time_col], errors="coerce")
    s = s.dropna()
    if s.empty:
        return {}
    return {
        "date_min": str(s.min().date()),
        "date_max": str(s.max().date()),
        "peak_hours": s.dt.hour.value_counts().head(6).to_dict(),
        "by_dow": s.dt.day_name().value_counts().to_dict(),
        "by_month": s.dt.to_period("M").astype(str).value_counts().head(12).sort_index().to_dict(),
    }

def build_data_report(df_all: pd.DataFrame, df_f: pd.DataFrame, filters: dict, groq_api_key: str = "") -> Dict[str, Any]:
    """
    Builds a full, data-driven report:
      - computed stats/tables
      - optional AI narrative using Groq (if key provided)
    """
    time_col = detect_time_col(df_f)
    lat_col, lon_col = detect_lat_lon(df_f)

    victims = pick_numeric_sum(df_f, ["victims", "total_victims", "injured_total", "injuries_total"]) or 0
    killed = pick_numeric_sum(df_f, ["killed", "deaths", "fatalities", "dead"])  # may be None
    serious = pick_numeric_sum(df_f, ["serious", "serious_injuries", "serious_injury", "injured_serious"])  # may be None

    summary = {
        "generated_at": datetime.datetime.now().isoformat(timespec="seconds"),
        "rows_filtered": int(len(df_f)),
        "rows_total": int(len(df_all)),
        "columns": int(df_all.shape[1]),
        "filters": {k: v for k, v in filters.items() if v},
        "victims_total": int(victims),
        "killed_total": int(killed) if isinstance(killed, int) else None,
        "serious_total": int(serious) if isinstance(serious, int) else None,
        "has_gps": bool(lat_col and lon_col),
        "gps_points": int(df_f[[lat_col, lon_col]].dropna().shape[0]) if (lat_col and lon_col) else 0,
        "time_column": time_col,
    }

    tables = {
        "top_provinces": _safe_value_counts(df_f, "location_province", 10),
        "top_locations": _safe_value_counts(df_f, "location_name", 12),
        "top_vehicle_type": _safe_value_counts(df_f, "vehicle_type", 10),
        "top_road_type": _safe_value_counts(df_f, "road_type", 10),
        "top_road_condition": _safe_value_counts(df_f, "road_condition", 10),
        "missingness_top": _missingness(df_f, 12),
    }

    time_stats = _time_insights(df_f, time_col) if time_col else {}
    insights = {
        "time": time_stats,
        "notes": [],
    }
    if not time_col:
        insights["notes"].append("No obvious date/time column detected; time-based trends were skipped.")
    if not (lat_col and lon_col):
        insights["notes"].append("No GPS coordinates detected; hotspot map uses province/location summaries.")

    # deterministic recommendations based on top patterns
    recs: List[str] = []
    if tables["top_provinces"]:
        top_prov = next(iter(tables["top_provinces"].keys()))
        recs.append(f"Target enforcement and awareness campaigns in {top_prov} (highest crash concentration in the filtered data).")
    if time_stats.get("peak_hours"):
        top_hour = next(iter(time_stats["peak_hours"].keys()))
        recs.append(f"Increase patrols and speed checks around {top_hour}:00 (peak crash hour).")
    if tables["top_vehicle_type"]:
        top_v = next(iter(tables["top_vehicle_type"].keys()))
        recs.append(f"Strengthen safety interventions for {top_v} (most frequent vehicle category in records).")
    if tables["top_road_condition"]:
        top_rc = next(iter(tables["top_road_condition"].keys()))
        recs.append(f"Prioritize fixes and warning signage where road condition is '{top_rc}'.")
    if not recs:
        recs = [
            "Improve speed management on high-risk corridors.",
            "Strengthen helmet/seatbelt checks and night visibility measures.",
            "Use hotspot analysis to target interventions with limited resources.",
        ]

    next_actions = [
        "Validate hotspots with police / RTDA records and field observations.",
        "Refresh the dataset weekly/monthly and re-run the report to track changes.",
        "Agree on 2–3 KPIs (fatalities, serious injuries, crash counts) for stakeholder reporting.",
    ]

    # Optional AI narrative (short, but grounded in computed stats)
    ai_narrative = None
    ai_recommendations = None
    if groq_api_key:
        try:
            llm = ChatGroq(
                model_name=os.getenv("GROQ_MODEL", "llama-3.3-70b-versatile"),
                api_key=groq_api_key,
                temperature=0.25,
            )
            stats_payload = {
                "summary": summary,
                "tables": {
                    "top_provinces": list(tables["top_provinces"].items())[:6],
                    "top_locations": list(tables["top_locations"].items())[:8],
                    "top_vehicle_type": list(tables["top_vehicle_type"].items())[:6],
                    "top_road_condition": list(tables["top_road_condition"].items())[:6],
                },
                "time": time_stats,
                "notes": insights["notes"],
            }

            prompt = PromptTemplate.from_template(
                "You are a Rwanda road safety analyst writing a professional stakeholder report.\n"
                "Use ONLY the provided stats. Do not invent numbers.\n\n"
                "STATS (JSON):\n{stats_json}\n\n"
                "Write:\n"
                "1) Executive summary (5–7 bullets)\n"
                "2) Key patterns (short paragraphs)\n"
                "3) 6 actionable recommendations (bullets, each linked to a pattern)\n"
            )
            chain = prompt | llm | StrOutputParser()
            txt = chain.invoke({"stats_json": json.dumps(stats_payload, ensure_ascii=False)})
            ai_narrative = txt

            # Try to extract recommendations section (best-effort)
            ai_recommendations = None
            if "recommendation" in txt.lower():
                ai_recommendations = txt
        except Exception as e:
            insights["notes"].append(f"AI narrative skipped (Groq error): {e}")

    report = {
        "title": "AutoRTC Road Safety Report",
        "summary": summary,
        "tables": tables,
        "insights": insights,
        "recommendations": recs,
        "next_actions": next_actions,
        "ai_narrative": ai_narrative,
    }
    return report


def report_to_pdf_bytes(report: Dict[str, Any]) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, leftMargin=2*cm, rightMargin=2*cm, topMargin=2*cm, bottomMargin=2*cm)
    styles = getSampleStyleSheet()
    story = []

    title = report.get("title", "AutoRTC Report")
    story.append(Paragraph(f"<b>{title}</b>", styles["Title"]))
    story.append(Spacer(1, 10))

    summary = report.get("summary", {})
    story.append(Paragraph("<b>Dataset overview</b>", styles["Heading2"]))
    ov = [
        ["Generated at", summary.get("generated_at", "")],
        ["Rows (filtered / total)", f"{summary.get('rows_filtered','-'):,} / {summary.get('rows_total','-'):,}"],
        ["Columns", str(summary.get("columns", "-"))],
        ["Victims total", str(summary.get("victims_total", "-"))],
        ["Killed total", str(summary.get("killed_total", "-"))],
        ["Serious total", str(summary.get("serious_total", "-"))],
        ["Time column", str(summary.get("time_column", "-"))],
        ["GPS points", str(summary.get("gps_points", "0"))],
    ]
    t = Table(ov, colWidths=[6*cm, 9*cm])
    t.setStyle(TableStyle([
        ("BACKGROUND", (0,0), (-1,0), colors.whitesmoke),
        ("BOX", (0,0), (-1,-1), 0.5, colors.lightgrey),
        ("INNERGRID", (0,0), (-1,-1), 0.25, colors.lightgrey),
        ("FONTNAME", (0,0), (-1,-1), "Helvetica"),
        ("FONTSIZE", (0,0), (-1,-1), 9),
        ("VALIGN", (0,0), (-1,-1), "TOP"),
    ]))
    story.append(t)
    story.append(Spacer(1, 12))

    # Key tables
    tables = report.get("tables", {})
    def add_kv_table(title_txt: str, data: Dict[str, int], max_rows: int = 10):
        story.append(Paragraph(f"<b>{title_txt}</b>", styles["Heading3"]))
        rows = [["Item", "Count"]]
        for k, v in list(data.items())[:max_rows]:
            rows.append([str(k), f"{int(v):,}"])
        tt = Table(rows, colWidths=[10*cm, 5*cm])
        tt.setStyle(TableStyle([
            ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#eef2ff")),
            ("TEXTCOLOR", (0,0), (-1,0), colors.HexColor("#0f172a")),
            ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"),
            ("BOX", (0,0), (-1,-1), 0.5, colors.lightgrey),
            ("INNERGRID", (0,0), (-1,-1), 0.25, colors.lightgrey),
            ("FONTSIZE", (0,0), (-1,-1), 9),
        ]))
        story.append(tt)
        story.append(Spacer(1, 10))

    if isinstance(tables.get("top_provinces"), dict) and tables["top_provinces"]:
        add_kv_table("Top provinces (by crash count)", tables["top_provinces"], 10)
    if isinstance(tables.get("top_locations"), dict) and tables["top_locations"]:
        add_kv_table("Top locations (by crash count)", tables["top_locations"], 10)
    if isinstance(tables.get("top_vehicle_type"), dict) and tables["top_vehicle_type"]:
        add_kv_table("Vehicle types (most frequent)", tables["top_vehicle_type"], 10)

    # Time insights (brief)
    time_stats = report.get("insights", {}).get("time", {}) or {}
    if time_stats:
        story.append(Paragraph("<b>Time patterns</b>", styles["Heading2"]))
        story.append(Paragraph(
            f"Date range: {time_stats.get('date_min','-')} to {time_stats.get('date_max','-')}.",
            styles["BodyText"]
        ))
        peak = time_stats.get("peak_hours", {}) or {}
        if peak:
            peak_txt = ", ".join([f"{h}:00 ({c})" for h, c in list(peak.items())[:6]])
            story.append(Paragraph(f"Peak hours: {peak_txt}.", styles["BodyText"]))
        story.append(Spacer(1, 10))

    # AI narrative
    ai_txt = report.get("ai_narrative")
    if ai_txt:
        story.append(PageBreak())
        story.append(Paragraph("<b>AI narrative (data-grounded)</b>", styles["Heading2"]))
        for para in str(ai_txt).split("\n"):
            para = para.strip()
            if para:
                story.append(Paragraph(para.replace("&", "&amp;"), styles["BodyText"]))
                story.append(Spacer(1, 6))

    # Recommendations
    story.append(Spacer(1, 8))
    story.append(Paragraph("<b>Recommendations</b>", styles["Heading2"]))
    for r in report.get("recommendations", [])[:12]:
        story.append(Paragraph(f"• {r}", styles["BodyText"]))
    story.append(Spacer(1, 8))
    story.append(Paragraph("<b>Next actions</b>", styles["Heading2"]))
    for a in report.get("next_actions", [])[:12]:
        story.append(Paragraph(f"• {a}", styles["BodyText"]))

    doc.build(story)
    buf.seek(0)
    return buf.getvalue()


def report_to_docx_bytes(report: Dict[str, Any]) -> bytes:
    doc = Document()
    doc.add_heading(report.get("title", "AutoRTC Report"), level=1)

    summary = report.get("summary", {})
    doc.add_paragraph(f"Generated at: {summary.get('generated_at','')}")
    doc.add_paragraph(f"Rows (filtered/total): {summary.get('rows_filtered','-')}/{summary.get('rows_total','-')}")
    doc.add_paragraph(f"Columns: {summary.get('columns','-')}")
    doc.add_paragraph(f"Victims total: {summary.get('victims_total','-')}")
    if summary.get("killed_total") is not None:
        doc.add_paragraph(f"Killed total: {summary.get('killed_total')}")
    if summary.get("serious_total") is not None:
        doc.add_paragraph(f"Serious total: {summary.get('serious_total')}")

    tables = report.get("tables", {})
    def add_table(title: str, data: Dict[str, int], n: int = 10):
        doc.add_heading(title, level=2)
        table = doc.add_table(rows=1, cols=2)
        hdr = table.rows[0].cells
        hdr[0].text = "Item"
        hdr[1].text = "Count"
        for k, v in list(data.items())[:n]:
            row_cells = table.add_row().cells
            row_cells[0].text = str(k)
            row_cells[1].text = str(int(v))
        doc.add_paragraph("")

    if isinstance(tables.get("top_provinces"), dict) and tables["top_provinces"]:
        add_table("Top provinces (by crash count)", tables["top_provinces"], 10)
    if isinstance(tables.get("top_locations"), dict) and tables["top_locations"]:
        add_table("Top locations (by crash count)", tables["top_locations"], 10)
    if isinstance(tables.get("top_vehicle_type"), dict) and tables["top_vehicle_type"]:
        add_table("Vehicle types (most frequent)", tables["top_vehicle_type"], 10)

    ai_txt = report.get("ai_narrative")
    if ai_txt:
        doc.add_heading("AI narrative (data-grounded)", level=2)
        for line in str(ai_txt).split("\n"):
            if line.strip():
                doc.add_paragraph(line.strip())

    doc.add_heading("Recommendations", level=2)
    for r in report.get("recommendations", []):
        doc.add_paragraph(str(r), style="List Bullet")

    doc.add_heading("Next actions", level=2)
    for a in report.get("next_actions", []):
        doc.add_paragraph(str(a), style="List Bullet")

    # Make text slightly bigger
    for p in doc.paragraphs:
        for run in p.runs:
            run.font.size = Pt(11)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def find_home_images(max_n: int = 3) -> List[str]:
    """
    Tries to auto-pick 3 images from the project folder.
    If you already have the 3 images used before, keep them in the same folder as main.py.
    """
    exts = (".png", ".jpg", ".jpeg", ".webp")
    candidates = []
    # prefer common names if they exist
    preferred = [
                "img1.jpeg", "img2.jpeg", "img3.jpeg",
    ]
    for name in preferred:
        if os.path.exists(name) and name.lower().endswith(exts):
            candidates.append(name)

    if len(candidates) < max_n:
        for fn in os.listdir("."):
            if fn.lower().endswith(exts) and fn not in candidates:
                # avoid logos if any
                if "logo" in fn.lower() or "icon" in fn.lower():
                    continue
                candidates.append(fn)
            if len(candidates) >= max_n:
                break
    return candidates[:max_n]


# ---------------- Sidebar ----------------
with st.sidebar:
    st.markdown("<h2 style='margin:0;color:#E91E63;'>💥🚗 AutoRTC AI</h2>", unsafe_allow_html=True)
    st.divider()
    c1, c2 = st.columns(2)
    c1.image("ur.png")
    c2.image("cmu.png")
    
    st.divider()
    pages = {"🏠 Home": "Home", "📊 Explore": "Explore", "📈 Models": "Models", "🤖 AI Hub": "AI Hub", "📄 Export": "Export"}
    labels = list(pages.keys())
    current_label = next((k for k, v in pages.items() if v == st.session_state.page), "🏠 Home")
    nav = st.radio("Navigation", labels, index=labels.index(current_label))
    st.session_state.page = pages[nav]
    st.divider()
    source = st.selectbox("Data Source", ["CSV", "JSON", "SQLite"], index=0)

    st.session_state.groq_api_key = st.text_input("Groq API Key", value=st.session_state.groq_api_key, type="password")


# ---------------- Data ----------------
df_raw = read_input_streamlit(st, source)
if df_raw is None:
    st.stop()
df = preprocess(df_raw)

with st.sidebar:
    st.markdown("---")
    prov_opts = sorted(df["location_province"].dropna().unique()) if "location_province" in df.columns else []
    location_province = st.multiselect("Province", prov_opts, default=prov_opts)

    loc_opts = sorted(df["location_name"].dropna().unique()) if "location_name" in df.columns else []
    location_name = st.multiselect("Location", loc_opts, default=loc_opts[:10])

filters = dict(location_province=location_province, location_name=location_name)
df_f = filter_df(df, filters)


# ---------------- Pages ----------------
def page_home():
    victims = pick_numeric_sum(df_f, ["victims", "total_victims", "injured_total", "injuries_total"]) or 0

    st.markdown(f"""
    <div class="hero">
      <h1 style="margin:0;">AutoRTC Intelligence</h1>
      <p style="color:var(--muted);font-weight:700;margin-top:8px;">
        Rwanda Road Safety Analytics • {len(df_f):,} crashes • {victims:,} victims
      </p>
    </div>
    """, unsafe_allow_html=True)

    # --- 3 images instead of map (as requested)
    imgs = find_home_images(3)
    st.markdown("<div class='section'><h3 style='margin-top:0;'>Overview</h3><p style='color:var(--muted);font-weight:700;margin-top:-6px;'>Quick visual context</p>", unsafe_allow_html=True)

    c1, c2, c3 = st.columns(3, gap="large")
    caps = [("Road safety", "Monitoring trends & hotspots"), ("Data insights", "Explore patterns by time & location"), ("Decision support", "Export stakeholder-ready reports")]
    cols = [c1, c2, c3]
    for i, col in enumerate(cols):
        with col:
            st.markdown("<div class='imgcard'>", unsafe_allow_html=True)
            if i < len(imgs) and os.path.exists(imgs[i]):
                st.image(imgs[i], use_container_width=True)
            else:
                # soft placeholder if images missing
                st.markdown(
                    "<div style='height:190px;background:linear-gradient(135deg,#eef7ff,#f2faf5);'></div>",
                    unsafe_allow_html=True
                )
            st.markdown(f"<div class='imgcap'>{caps[i][0]}</div>", unsafe_allow_html=True)
            st.markdown(f"<div class='imgsub'>{caps[i][1]}</div>", unsafe_allow_html=True)
            st.markdown("</div>", unsafe_allow_html=True)

    st.markdown("</div>", unsafe_allow_html=True)

    # Quick Launch
    st.markdown("<div class='section'><h3 style='margin-top:0;'>Quick Launch</h3>", unsafe_allow_html=True)
    c1, c2, c3, c4 = st.columns(4)

    with c1:
        st.markdown('<div class="launch card-blue"><div class="launch-title">📊 Explore</div><div class="launch-desc">Trends & hotspots</div></div>', unsafe_allow_html=True)
        st.markdown('<div class="btn-blue">', unsafe_allow_html=True)
        if st.button("Open Explore", key="btn_home_explore"):
            st.session_state.page = "Explore"; st.rerun()
        st.markdown("</div>", unsafe_allow_html=True)

    with c2:
        st.markdown('<div class="launch card-yellow"><div class="launch-title">📈 Models</div><div class="launch-desc">Predictive risk</div></div>', unsafe_allow_html=True)
        st.markdown('<div class="btn-yellow">', unsafe_allow_html=True)
        if st.button("Open Models", key="btn_home_models"):
            st.session_state.page = "Models"; st.rerun()
        st.markdown("</div>", unsafe_allow_html=True)

    with c3:
        st.markdown('<div class="launch card-green"><div class="launch-title">🤖 AI Hub</div><div class="launch-desc">Full report + chat</div></div>', unsafe_allow_html=True)
        st.markdown('<div class="btn-green">', unsafe_allow_html=True)
        if st.button("Open AI Hub", key="btn_home_ai"):
            st.session_state.page = "AI Hub"; st.rerun()
        st.markdown("</div>", unsafe_allow_html=True)

    with c4:
        st.markdown('<div class="launch card-dark"><div class="launch-title">📄 Export</div><div class="launch-desc">PDF • Word • JSON</div></div>', unsafe_allow_html=True)
        st.markdown('<div class="btn-dark">', unsafe_allow_html=True)
        if st.button("Open Export", key="btn_home_export"):
            st.session_state.page = "Export"; st.rerun()
        st.markdown("</div>", unsafe_allow_html=True)

    st.markdown("</div>", unsafe_allow_html=True)


def page_ai():
    st.markdown("<div class='hero'><h2 style='margin:0;'>🤖 AI Hub</h2><p style='color:var(--muted);font-weight:700;margin-top:6px;'>Generate a full, data-driven report and chat with your filtered dataset.</p></div>", unsafe_allow_html=True)

    left, right = st.columns([1.0, 1.0], gap="large")

    # --- Full report generator
    with left:
        st.markdown('<div class="ai-agent">', unsafe_allow_html=True)
        st.markdown("<h3 style='margin-top:0;'>✨ Report Generator</h3>", unsafe_allow_html=True)

        st.markdown('<div class="btn-blue">', unsafe_allow_html=True)
        if st.button("Generate Full Report", key="btn_generate_report", use_container_width=True):
            st.session_state.latest_report = build_data_report(df, df_f, filters, groq_api_key=st.session_state.groq_api_key)
            st.success("Report generated. Open Export to download.")
        st.markdown("</div>", unsafe_allow_html=True)

        if st.session_state.latest_report:
            # Clean preview: show key metrics, not raw JSON dump
            rep = st.session_state.latest_report
            s = rep.get("summary", {})
            st.markdown(
                f"**Rows (filtered/total):** {s.get('rows_filtered','-'):,} / {s.get('rows_total','-'):,}  \n"
                f"**Columns:** {s.get('columns','-')}  \n"
                f"**Victims:** {s.get('victims_total','-')}  \n"
                f"**Time column:** {s.get('time_column','-')}  \n"
                f"**GPS points:** {s.get('gps_points','0')}"
            )
            with st.expander("Preview full report (JSON)"):
                st.json(rep)
        else:
            st.info("Click **Generate Full Report** to create a detailed stakeholder report from your uploaded data.")

        st.markdown("</div>", unsafe_allow_html=True)

    # --- Chat (Groq)
    with right:
        st.markdown('<div class="ai-chat">', unsafe_allow_html=True)
        st.markdown("<h3 style='margin-top:0;'>💬 Data Chat</h3>", unsafe_allow_html=True)

        if not st.session_state.groq_api_key:
            st.warning("Add your Groq API key in the sidebar to enable chat.")
            st.markdown("</div>", unsafe_allow_html=True)
            return

        # lightweight context
        cols = ", ".join([f"{c} ({str(df_f[c].dtype)})" for c in df_f.columns[:35]])
        head = df_f.head(6).to_dict(orient="records")
        data_ctx = f"Rows filtered: {len(df_f)}\nColumns: {cols}\nSample rows: {head}"

        llm = ChatGroq(
            model_name=os.getenv("GROQ_MODEL", "llama-3.3-70b-versatile"),
            api_key=st.session_state.groq_api_key,
            temperature=0.3,
        )
        prompt = PromptTemplate.from_template(
            "You are a Rwanda road safety analyst. Be short, practical, and beginner-friendly.\n\n"
            "DATA CONTEXT:\n{data_context}\n\n"
            "CHAT HISTORY:\n{history}\n\n"
            "QUESTION:\n{question}\n"
        )
        chain = prompt | llm | StrOutputParser()

        sid = "autorrc_chat"
        if sid not in st.session_state:
            st.session_state[sid] = []
        # show last messages stored in session (simple + robust)
        for m in st.session_state[sid][-10:]:
            with st.chat_message(m["role"]):
                st.write(m["content"])

        user_msg = st.chat_input("Ask about hotspots, provinces, peak hours…")
        if user_msg:
            st.session_state[sid].append({"role": "user", "content": user_msg})
            with st.chat_message("user"):
                st.write(user_msg)

            try:
                # We include only last few messages as plain text history
                hist_txt = "\n".join([f'{m["role"]}: {m["content"]}' for m in st.session_state[sid][-8:]])
                ans = chain.invoke({"question": user_msg, "data_context": data_ctx, "history": hist_txt})
            except Exception as e:
                ans = f"Chat error: {e}"

            st.session_state[sid].append({"role": "assistant", "content": ans})
            with st.chat_message("assistant"):
                st.write(ans)

        st.markdown("</div>", unsafe_allow_html=True)


def page_export():
    st.markdown("<div class='hero'><h2 style='margin:0;'>📄 Export</h2><p style='color:var(--muted);font-weight:700;margin-top:6px;'>Download a full report (JSON, PDF, Word).</p></div>", unsafe_allow_html=True)

    if not st.session_state.latest_report:
        st.info("Generate a **Full Report** in AI Hub first.")
        return

    report = st.session_state.latest_report
    json_bytes = json.dumps(report, indent=2, ensure_ascii=False).encode("utf-8")

    st.markdown("<div class='section'><h3 style='margin-top:0;'>Downloads</h3>", unsafe_allow_html=True)
    b1, b2, b3 = st.columns(3, gap="large")

    with b1:
        st.markdown('<div class="json-btn">', unsafe_allow_html=True)
        st.download_button("🧾 JSON (Full)", data=json_bytes, file_name="autorrc_report.json", mime="application/json", use_container_width=True, key="dl_json")
        st.markdown("</div>", unsafe_allow_html=True)

    with b2:
        st.markdown('<div class="pdf-btn">', unsafe_allow_html=True)
        st.download_button("📕 PDF (Full)", data=report_to_pdf_bytes(report), file_name="autorrc_report.pdf", mime="application/pdf", use_container_width=True, key="dl_pdf")
        st.markdown("</div>", unsafe_allow_html=True)

    with b3:
        st.markdown('<div class="word-btn">', unsafe_allow_html=True)
        st.download_button("📝 Word (Full)", data=report_to_docx_bytes(report), file_name="autorrc_report.docx",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True, key="dl_docx")
        st.markdown("</div>", unsafe_allow_html=True)

    with st.expander("Preview report"):
        st.json(report)

    st.markdown("</div>", unsafe_allow_html=True)


# ---------------- Router ----------------
if st.session_state.page == "Home":
    page_home()
elif st.session_state.page == "Explore":
    descriptive.render(df, filters)
elif st.session_state.page == "Models":
    predictive.render(df, filters)
elif st.session_state.page == "AI Hub":
    page_ai()
else:
    page_export()
